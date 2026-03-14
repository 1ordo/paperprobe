import logging
import re
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

SECTION_PATTERNS = [
    (r"(?i)^abstract\b", "abstract"),
    (r"(?i)^introduction\b", "introduction"),
    (r"(?i)^background\b", "introduction"),
    (r"(?i)^methods?\b", "methods"),
    (r"(?i)^materials?\s+and\s+methods?\b", "methods"),
    (r"(?i)^results?\b", "results"),
    (r"(?i)^findings?\b", "results"),
    (r"(?i)^discussion\b", "discussion"),
    (r"(?i)^conclusions?\b", "discussion"),
    (r"(?i)^limitations?\b", "discussion"),
    (r"(?i)^references?\b", "references"),
    (r"(?i)^bibliography\b", "references"),
    (r"(?i)^acknowledg", "other"),
    (r"(?i)^appendix", "other"),
    (r"(?i)^supplementary", "other"),
    (r"(?i)^table\s+\d", "table"),
    (r"(?i)^figure\s+\d", "figure"),
]


@dataclass
class ParsedSection:
    section_type: str
    heading: str
    content: str
    page_start: int
    page_end: int
    position_order: int


@dataclass
class ParsedDocument:
    title: str | None = None
    authors: str | None = None
    year: int | None = None
    page_count: int = 0
    full_text: str = ""
    sections: list[ParsedSection] = field(default_factory=list)


def classify_section(heading: str) -> str:
    heading_clean = heading.strip()
    for pattern, section_type in SECTION_PATTERNS:
        if re.match(pattern, heading_clean):
            return section_type
    return "other"


def parse_pdf(file_path: str) -> ParsedDocument:
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    result = ParsedDocument(page_count=len(doc))

    all_text_parts = []
    page_texts = {}

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")
        page_texts[page_num + 1] = text
        all_text_parts.append(text)

    result.full_text = "\n\n".join(all_text_parts)

    # Try to extract title from first page
    if page_texts.get(1):
        lines = page_texts[1].strip().split("\n")
        non_empty = [l.strip() for l in lines if l.strip()]
        if non_empty:
            result.title = non_empty[0][:500]

    # Build sections by detecting headings
    current_section = None
    current_content_parts = []
    current_page_start = 1
    section_order = 0

    for page_num, text in sorted(page_texts.items()):
        lines = text.split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            section_type = classify_section(stripped)
            is_heading = (
                section_type != "other"
                or (len(stripped) < 80 and stripped.isupper())
                or (len(stripped) < 80 and re.match(r"^\d+\.?\s+[A-Z]", stripped))
            )

            if is_heading and section_type != "other":
                # Save previous section
                if current_section and current_content_parts:
                    result.sections.append(ParsedSection(
                        section_type=current_section,
                        heading=current_section,
                        content="\n".join(current_content_parts),
                        page_start=current_page_start,
                        page_end=page_num,
                        position_order=section_order,
                    ))
                    section_order += 1

                current_section = section_type
                current_content_parts = [stripped]
                current_page_start = page_num
            else:
                current_content_parts.append(stripped)

    # Save last section
    if current_section and current_content_parts:
        result.sections.append(ParsedSection(
            section_type=current_section,
            heading=current_section,
            content="\n".join(current_content_parts),
            page_start=current_page_start,
            page_end=result.page_count,
            position_order=section_order,
        ))

    # If no sections detected, create one big section
    if not result.sections and result.full_text.strip():
        result.sections.append(ParsedSection(
            section_type="full_text",
            heading="Full Document",
            content=result.full_text,
            page_start=1,
            page_end=result.page_count,
            position_order=0,
        ))

    doc.close()
    return result


def parse_docx(file_path: str) -> ParsedDocument:
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    result = ParsedDocument()

    all_text_parts = []
    current_section = None
    current_content_parts = []
    section_order = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        all_text_parts.append(text)
        is_heading = para.style.name.startswith("Heading")
        section_type = classify_section(text) if is_heading else "other"

        if is_heading and section_type != "other":
            if current_section and current_content_parts:
                result.sections.append(ParsedSection(
                    section_type=current_section,
                    heading=current_section,
                    content="\n".join(current_content_parts),
                    page_start=1,
                    page_end=1,
                    position_order=section_order,
                ))
                section_order += 1
            current_section = section_type
            current_content_parts = [text]
        else:
            current_content_parts.append(text)

    if current_section and current_content_parts:
        result.sections.append(ParsedSection(
            section_type=current_section,
            heading=current_section,
            content="\n".join(current_content_parts),
            page_start=1,
            page_end=1,
            position_order=section_order,
        ))

    result.full_text = "\n\n".join(all_text_parts)
    if not result.title and all_text_parts:
        result.title = all_text_parts[0][:500]

    if not result.sections and result.full_text.strip():
        result.sections.append(ParsedSection(
            section_type="full_text",
            heading="Full Document",
            content=result.full_text,
            page_start=1,
            page_end=1,
            position_order=0,
        ))

    return result


def parse_document(file_path: str) -> ParsedDocument:
    if file_path.lower().endswith(".pdf"):
        return parse_pdf(file_path)
    elif file_path.lower().endswith(".docx"):
        return parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_path}")


def chunk_text(text: str, chunk_size: int = 512, overlap: int = 50) -> list[dict]:
    """Split text into overlapping chunks by word count."""
    words = text.split()
    chunks = []
    start = 0
    char_pos = 0

    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunk_words = words[start:end]
        chunk_text = " ".join(chunk_words)

        # Calculate character positions
        chunk_start = text.find(chunk_words[0], char_pos) if chunk_words else char_pos
        chunk_end = chunk_start + len(chunk_text)

        chunks.append({
            "text": chunk_text,
            "char_start": chunk_start,
            "char_end": chunk_end,
            "word_start": start,
            "word_end": end,
        })

        if end >= len(words):
            break
        start = end - overlap
        char_pos = chunk_start

    return chunks
